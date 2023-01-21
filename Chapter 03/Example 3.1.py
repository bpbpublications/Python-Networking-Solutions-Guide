import docx				# Import python-docx module		
document = docx.Document()		# Call document function to call other functions

document.add_heading('PYTHON COURSE V1.0', 0)	    # Add heading to word document
						
p = document.add_paragraph('We are learning ')    # Create a new paragraph
p.add_run('Python. ').bold = True		    # Add characters in bold
p.add_run('for ')					    # Add characters in default style
p.add_run('Network Automation.').italic = True	    # Add characters in italic

# Add 2 lines of bullet style text
document.add_paragraph('Lesson-1 Introduction', style='List Bullet')
document.add_paragraph('Lesson-2 Installation', style='List Bullet')

# Add 2 lines of Numbered list
document.add_paragraph("What is Python?", style='List Number')
document.add_paragraph("How to install Python?", style='List Number')

document.add_picture('logo.jpg',  width=docx.shared.Inches(2))   # Add Picture

document.add_heading('TABLE-1', 2)            # Add Heading with size "2"
table = document.add_table(rows=2, cols=2)    # Add Table with 2 rows and 2 columns
table.style = document.styles['Table Grid']   
cell = table.cell(0, 0)				# Fill Table by cells
cell.text = "python"
cell = table.cell(0, 1)
cell.text = "automation"

row = table.rows[1]				  # Fill Table by cells in alternative way
row.cells[0].text = 'network'
row.cells[1].text = 'engineers'

row = table.add_row()				# Add new row to table

document.save('test.docx')				# Save all changes to excel file